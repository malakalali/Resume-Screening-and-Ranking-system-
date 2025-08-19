import os
import PyPDF2
import docx
import re
from typing import Dict, List, Optional, Tuple

class DocumentParser:
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.doc'}
    
    def parse_document(self, file_path: str) -> Tuple[bool, str]:
        if not os.path.exists(file_path):
            return False, "File does not exist"
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension not in self.supported_extensions:
            return False, f"Unsupported file format: {file_extension}"
        
        try:
            if file_extension == '.pdf':
                return self._parse_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._parse_docx(file_path)
        except Exception as e:
            return False, f"Error parsing document: {str(e)}"
    
    def _parse_pdf(self, file_path: str) -> Tuple[bool, str]:
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if len(pdf_reader.pages) == 0:
                    return False, "PDF file is empty"
                
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
                
                if not text_content:
                    return False, "No text content found in PDF"
                
                full_text = '\n'.join(text_content)
                cleaned_text = self._clean_text(full_text)
                
                return True, cleaned_text
                
        except Exception as e:
            return False, f"PDF parsing error: {str(e)}"
    
    def _parse_docx(self, file_path: str) -> Tuple[bool, str]:
        try:
            doc = docx.Document(file_path)
            
            if not doc.paragraphs:
                return False, "DOCX file is empty"
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            if not text_content:
                return False, "No text content found in DOCX"
            
            full_text = '\n'.join(text_content)
            cleaned_text = self._clean_text(full_text)
            
            return True, cleaned_text
            
        except Exception as e:
            return False, f"DOCX parsing error: {str(e)}"
    
    def _clean_text(self, text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)]', '', text)
        text = text.strip()
        return text
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        sections = {
            'full_text': text,
            'summary': '',
            'experience': '',
            'education': '',
            'skills': ''
        }
        
        lines = text.split('\n')
        current_section = 'summary'
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            line_lower = line.lower()
            
            if any(keyword in line_lower for keyword in ['experience', 'work history', 'employment']):
                current_section = 'experience'
            elif any(keyword in line_lower for keyword in ['education', 'academic', 'degree']):
                current_section = 'education'
            elif any(keyword in line_lower for keyword in ['skills', 'technologies', 'programming']):
                current_section = 'skills'
            
            if current_section in sections:
                sections[current_section] += line + '\n'
        
        for key in sections:
            sections[key] = sections[key].strip()
        
        return sections 